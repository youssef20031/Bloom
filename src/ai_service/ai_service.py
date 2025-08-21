import os
from flask import Flask, request, jsonify, Response, send_file
from flask_cors import CORS
from langchain_community.document_loaders import PyPDFLoader, DirectoryLoader
from langchain_community.vectorstores import Chroma
from langchain_core.messages import HumanMessage, AIMessage
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain.chains import create_history_aware_retriever, create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from io import BytesIO
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.units import inch
from reportlab.lib import colors


# Initialize Flask App
app = Flask(__name__)
CORS(app) # Enable Cross-Origin Resource Sharing

# --- GLOBAL VARIABLES & ONE-TIME SETUP ---
vector_store = None
rag_chain = None
llm = None

def initialize_ai_components():
    """Loads documents, creates vector store, and builds the RAG chain."""
    global vector_store, rag_chain, llm

    print("--- Initializing AI Components ---")

    if 'GOOGLE_API_KEY' not in os.environ:
        print("ERROR: GOOGLE_API_KEY environment variable not set.")
        return

    print("Loading and processing PDF documents...")
    loader = DirectoryLoader('./', glob="**/*.pdf", loader_cls=PyPDFLoader, show_progress=True, use_multithreading=True)
    
    all_docs = loader.load()
    
    if not all_docs:
        print("\nERROR: No PDF files found in the current directory.")
        return

    print(f"Successfully loaded {len(all_docs)} pages from all PDF documents.")
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=200)
    doc_splits = text_splitter.split_documents(all_docs)

    print(f"Split documents into {len(doc_splits)} chunks.")
    print("Creating vector store with Gemini embeddings...")
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    vector_store = Chroma.from_documents(documents=doc_splits, embedding=embeddings)
    retriever = vector_store.as_retriever(search_kwargs={"k": 10})
    print("Vector store and retriever are ready.")

    print("Building the conversational RAG chain...")
    llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.3) # Updated model

    contextualize_q_system_prompt = (
        "Given a chat history and the latest user question "
        "formulate a standalone question which can be understood "
        "without the chat history. Do NOT answer the question, "
        "just reformulate it if needed and otherwise return it as is."
    )
    contextualize_q_prompt = ChatPromptTemplate.from_messages([
        ("system", contextualize_q_system_prompt),
        MessagesPlaceholder("chat_history"),
        ("human", "{input}"),
    ])
    history_aware_retriever = create_history_aware_retriever(llm, retriever, contextualize_q_prompt)

    qa_system_prompt = (
    "You are an expert-level Dell Technical Presales Assistant, acting as a proactive partner to a Dell presales engineer on a live customer call. "
    "Your knowledge base is a comprehensive set of Dell's technical documents for networking, storage, and servers. "
    "Your primary goals are:\n"
    "1.  **Provide Proactive Recommendations:** When the engineer provides a use case (e.g., 'HPC compute node', 'VDI storage'), immediately suggest a strong, common Dell product from your documents that fits the description. Do not just ask for more information initially.\n"
    "2.  **Justify Your Suggestion:** Briefly explain *why* your recommendation is a good fit, connecting 2-3 key technical specs from the documents to the stated requirements (e.g., 'The R660xs is ideal for density because it is a 1U server...').\n"
    "3.  **Offer Relevant Alternatives:** If applicable, suggest a clear alternative (e.g., a 2U option for more expansion, a different series for a different feature set) to show breadth of knowledge.\n"
    "4.  **Ask High-Impact Clarifying Questions:** After providing your initial recommendation, ask one or two targeted questions that will best help refine the solution (e.g., 'Is GPU acceleration required?', 'What level of data protection is needed?'). Avoid long, generic lists of questions.\n"
    "5.  **Be Precise and Factual:** For direct questions ('What are the dimensions of X?'), provide the exact data from the context. You MUST base your answers strictly on the provided documents. If the information is not present, state: 'That specific information is not available in my current documents.' Do not invent information.\n\n"
    "Your tone should be expert, concise, and collaborative. "
    "Context Documents:\n{context}"
    )
    qa_prompt = ChatPromptTemplate.from_messages([
        ("system", qa_system_prompt),
        MessagesPlaceholder("chat_history"),
        ("human", "{input}"),
    ])

    question_answer_chain = create_stuff_documents_chain(llm, qa_prompt)
    rag_chain = create_retrieval_chain(history_aware_retriever, question_answer_chain)
    
    print("--- AI Components Initialized Successfully ---")

# --- API ENDPOINTS ---

@app.route('/chat', methods=['POST'])
def chat():
    if not rag_chain:
        return jsonify({"error": "RAG chain not initialized"}), 500
    data = request.json
    user_input = data.get('input')
    chat_history_json = data.get('chat_history', [])

    # --- FIX #1: Correctly parse the list of lists format ---
    # msg[0] is the sender ('human' or 'ai'), msg[1] is the content
    chat_history = [
        HumanMessage(content=msg[1]) if msg[0] == 'human' else AIMessage(content=msg[1])
        for msg in chat_history_json
    ]

    def generate():
        for chunk in rag_chain.stream({"input": user_input, "chat_history": chat_history}):
            if "answer" in chunk:
                yield chunk['answer']
    return Response(generate(), mimetype='text/plain')

def _build_report_prompt_and_chain(report_type: str):
    """Create the prompt+chain for either 'human' or 'json' report types."""
    if report_type == 'human':
        prompt_template = (
            "You are a presales analyst. Generate a structured summary report "
            "based on the provided presales call transcript. Extract the key customer requirements, products discussed, and actionable next steps. "
            "The report must be formatted exactly as follows. If a section has no information, write 'Not specified'.\n\n"
            "**Presales Call Summary Report:**\n\n"
            "**1. Identified Customer Needs & Environment:**\n"
            "- **Primary Workload(s):** (e.g., VDI, OLTP Database, SD-WAN, Data Analytics)\n"
            "- **Key Performance Requirements:** (e.g., Low Latency, All-Flash, High IOPS, Specific bandwidth)\n"
            "- **Capacity & Scalability Needs:** (e.g., Started at 150TB, needs to scale)\n"
            "- **Connectivity Requirements:** (e.g., 10GbE SFP+, 400GbE, Redundant power)\n"
            "- **Must-Have Features:** (e.g., Data reduction, Disaster Recovery, Multi-tenancy)\n\n"
            "**2. Dell Products Discussed / Recommended:**\n"
            "- **Product Family/Model(s):** (e.g., PowerSwitch Z-series, PowerStore 5200T, VEP4600)\n"
            "- **Key Specifications Mentioned:** (List any specific ports, speeds, or capacities that were discussed)\n"
            "- **Reason for Recommendation:** (Briefly explain why this solution fits the customer's requirements based on the conversation)\n\n"
            "**3. Outstanding Questions & Information Gaps:**\n"
            "- (List any information the presales engineer still needs from the customer)\n\n"
            "**4. Actionable Next Steps:**\n"
            "- [ ] Follow up on outstanding questions.\n"
            "- [ ] Prepare a detailed technical proposal and quote for the recommended solution.\n\n"
            "--- CONVERSATION TRANSCRIPT ---\n{history}"
        )
    else:  # JSON report
        prompt_template = (
            "You are a data extraction bot. Analyze the conversation transcript and "
            "extract key requirements into a structured JSON format. "
            "Provide ONLY a valid JSON object as your final output. Do not add any explanatory text before or after the JSON. "
            "If a value isn't mentioned, use 'Not Specified'.\n\n"
            "JSON structure:\n"
            '{{"workload_types": [], "capacity_needs": "", "performance_needs": [], '
            '"connectivity_needs": [], "key_features_requested": [], "recommended_products": []}}\n\n'
            "TRANSCRIPT:\n{history}"
        )

    report_prompt = ChatPromptTemplate.from_template(prompt_template)
    report_chain = report_prompt | llm
    return report_chain


def _format_chat_history_as_text(chat_history):
    # --- FIX #2: Correctly parse the list of lists format here as well ---
    return "\n".join([f"{'Presales Engineer' if msg[0] == 'human' else 'Assistant'}: {msg[1]}" for msg in chat_history])


def generate_report(chat_history, report_type):
    report_chain = _build_report_prompt_and_chain(report_type)
    history_str = _format_chat_history_as_text(chat_history)

    def generate():
        for chunk in report_chain.stream({"history": history_str}):
            yield chunk.content
    
    return Response(generate(), mimetype='text/plain')


def generate_report_text(chat_history, report_type):
    report_chain = _build_report_prompt_and_chain(report_type)
    history_str = _format_chat_history_as_text(chat_history)
    result = report_chain.invoke({"history": history_str})
    content = getattr(result, 'content', None)
    return content if content is not None else str(result)

# ... (The rest of the file for building DOCX and PDF remains the same) ...

def build_docx_from_report(report_text: str) -> BytesIO:
    document = Document()
    document.sections[0].left_margin = Inches(1)
    document.sections[0].right_margin = Inches(1)
    
    lines = report_text.splitlines()
    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue

        if line.startswith("**") and line.endswith(":**"):
            heading_text = re.sub(r"^\*\*(.*?):\*\*$", r"\1", line)
            p = document.add_paragraph()
            p.add_run(heading_text).bold = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            continue
        
        if line.startswith("- **") and line.endswith(":**"):
            p = document.add_paragraph(style='List Bullet')
            sub_heading_text = re.search(r'\*\*(.*?):\*\*', line)
            if sub_heading_text:
                p.add_run(f"{sub_heading_text.group(1)}:").bold = True
                rest_of_line = line.split(':**', 1)[1].strip()
                if rest_of_line:
                    p.add_run(f" {rest_of_line}")
            else:
                 p.add_run(line[2:])
            continue

        if line.startswith("- "):
            p = document.add_paragraph(line[2:], style='List Bullet')
            continue

        document.add_paragraph(line)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def build_pdf_from_report(report_text: str) -> BytesIO:

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=inch*0.75, leftMargin=inch*0.75, topMargin=inch*0.75, bottomMargin=inch*0.75)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='SectionHeading', parent=styles['h2'], spaceBefore=12, spaceAfter=6, textColor=colors.HexColor('#007DB8')))
    styles.add(ParagraphStyle(name='Body', parent=styles['BodyText'], fontSize=10, leading=14))
    styles.add(ParagraphStyle(name='ListItem', parent=styles['Body'], leftIndent=18))
    
    flow = []
    flow.append(Paragraph('Presales Call Summary Report', styles['h1']))
    flow.append(Spacer(1, 0.2 * inch))

    lines = report_text.splitlines()
    current_list_items = []

    def flush_list():
        nonlocal current_list_items
        if current_list_items:
            list_items_for_flowable = []
            for item_text in current_list_items:
                 formatted_text = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", item_text)
                 list_items_for_flowable.append(ListItem(Paragraph(formatted_text, styles['ListItem'])))

            flow.append(ListFlowable(list_items_for_flowable, bulletType='bullet', start='-'))
            current_list_items = []

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue

        if line.startswith("**") and line.endswith(":**"):
            flush_list()
            heading_text = re.sub(r"^\*\*(.*?):\*\*$", r"\1", line)
            flow.append(Paragraph(heading_text, styles['SectionHeading']))
            continue
        
        if line.startswith("- "):
             current_list_items.append(line[2:])
             continue
        
        flush_list()
        flow.append(Paragraph(line, styles['Body']))
    
    flush_list() 
    doc.build(flow)
    buffer.seek(0)
    return buffer


@app.route('/report', methods=['POST'])
def get_human_report():
    chat_history = request.json.get('chat_history', [])
    return generate_report(chat_history, 'human')

@app.route('/json_report', methods=['POST'])
def get_json_report():
    chat_history = request.json.get('chat_history', [])
    return generate_report(chat_history, 'json')

@app.route('/report_docx', methods=['POST'])
def get_human_report_docx():
    chat_history = request.json.get('chat_history', [])
    report_text = generate_report_text(chat_history, 'human')
    docx_buffer = build_docx_from_report(report_text)
    return send_file(
        docx_buffer,
        as_attachment=True,
        download_name='dell_presales_report.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/report_pdf', methods=['POST'])
def get_human_report_pdf():
    chat_history = request.json.get('chat_history', [])
    report_text = generate_report_text(chat_history, 'human')
    pdf_buffer = build_pdf_from_report(report_text)
    return send_file(
        pdf_buffer,
        as_attachment=True,
        download_name='dell_presales_report.pdf',
        mimetype='application/pdf'
    )


if __name__ == '__main__':
    initialize_ai_components()
    if rag_chain:
        print("Starting Flask server...")
        app.run(host='0.0.0.0', port=5001, debug=False)
    else:
        print("Could not start Flask server because AI components failed to initialize.")